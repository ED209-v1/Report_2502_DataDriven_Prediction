# -*- coding: utf-8 -*-
"""
파일명  : 시간vs순번_회귀비교_소스코드_상세주석.py
작성자  : (작성일 자동)
목적    : 동일한 진동 센서 데이터셋에 대해
          1) "시간 기반(time_index)" 과 2) "자동 순번(NUM)" 두 방식으로
          단순회귀, 다중회귀(시차 변수 포함), KNN 회귀를 적용하고
          성능지표(MAE, RMSE, R^2)를 비교하며, 결과 그래프와 Word 리포트를 생성한다.

사용 데이터: mixing_actuator.csv
필수 컬럼 : Date, Sensor (Quality가 있어도 무방하나 본 분석에서는 회귀 대상이 아니므로 미사용)
출력물    :
  - regression_time_based.png  : 시간 기반 예측 비교 그래프
  - regression_num_based.png   : 순번 기반 예측 비교 그래프
  - 시간vs순번_회귀비교_리포트.docx : 두 방식의 성능 비교 Word 리포트

실행 환경 참고:
  - matplotlib는 기본 설정(스타일 미지정)을 사용 (요구사항)
  - 각 차트는 단일 플롯으로 생성 (요구사항)
"""

# =============== [1] 라이브러리 임포트 ===============
# 데이터 처리
import pandas as pd              # 표 형식 데이터 처리(DataFrame)
import numpy as np               # 수치 연산(배열, 수학 함수)

# 시각화
import matplotlib.pyplot as plt  # 그래프 생성(기본 스타일 사용)

# 학습/검증 데이터 분할
from sklearn.model_selection import train_test_split

# 회귀 모델
from sklearn.linear_model import LinearRegression      # 선형 회귀(단순/다중에 모두 사용)
from sklearn.neighbors import KNeighborsRegressor      # KNN 회귀

# 성능 지표
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score

# 리포트(Word) 생성
from docx import Document
from docx.shared import Inches

# =============== [2] 사용자 설정 ===============
# CSV 파일 경로: 동일 폴더에 있을 경우 파일명만 기입
CSV_PATH = "mixing_actuator.csv"

# 출력 파일 경로(동일 폴더 저장)
PLOT_TIME_PATH = "regression_time_based.png"     # 시간 기반 비교 그래프
PLOT_NUM_PATH  = "regression_num_based.png"      # 순번 기반 비교 그래프
DOCX_PATH      = "시간vs순번_회귀비교_리포트.docx"  # Word 리포트 파일명

# =============== [3] 데이터 로드 및 전처리 ===============
# 3-1) CSV 로드: pandas의 read_csv로 데이터프레임 생성
df = pd.read_csv(CSV_PATH)

# 3-2) Date 컬럼을 문자열 -> datetime으로 변환 (시간 연산을 위해 필수)
df["Date"] = pd.to_datetime(df["Date"])

# 3-3) 시간순으로 정렬 (시계열 일관성을 보장; 분석 및 시차 생성에 중요)
df = df.sort_values(by="Date").reset_index(drop=True)

# 3-4) 회귀 타깃인 Sensor에서 결측 제거 (결측이 있으면 학습 불가)
df = df.dropna(subset=["Sensor"])

# =============== [4] 파생 변수 생성(시간/순번/시차) ===============
# 4-1) 시간 기반 독립변수(time_index) 생성
#      - 기준시각(df["Date"].min())부터 각 관측시점까지의 "초 단위 경과시간"
#      - 실수 스케일이며 수집간격 불균등이 반영됨
df["time_index"] = (df["Date"] - df["Date"].min()).dt.total_seconds()

# 4-2) 자동 순번(NUM) 생성
#      - 0,1,2,... 순차 인덱스
#      - 관측샘플 간 간격이 균등하다고 가정하는 효과
df["NUM"] = np.arange(len(df))

# 주의) 아래의 시차 변수는 분석 함수 내부에서 각 방식별로 다시 생성하므로
#      여기서는 만들지 않지만, 참고로 기술함:
# df["Sensor_t-1"] = df["Sensor"].shift(1)
# df["Sensor_t-2"] = df["Sensor"].shift(2)

# =============== [5] 공통 회귀 분석 함수 정의 ===============
def run_regression_pipeline(X: pd.DataFrame, y: pd.Series):
    """
    목적 : 주어진 X, y에 대해 (1) 단순회귀, (2) 다중회귀(시차 포함), (3) KNN 회귀 수행
    입력 :
      - X : 독립변수 데이터프레임 (한 개 컬럼: time_index 또는 NUM)
      - y : 종속변수 시리즈 (Sensor)
    출력 :
      - dict 타입 결과 객체
        {
          "linear":   {"y_test":..., "y_pred":..., "MAE":..., "RMSE":..., "R2":...},
          "multiple": {...},
          "knn":      {...}
        }
    주석 :
      - 단순회귀는 X 한 컬럼만 사용
      - 다중/knn은 X + (Sensor_t-1, Sensor_t-2) 2개 시차 특성을 추가한 3개 컬럼 사용
    """
    results = {}  # 모든 모델 결과를 담을 딕셔너리

    # ---------- (A) 단순회귀 ----------
    # 데이터 분할 (학습:검증 = 8:2, 난수고정으로 재현성 보장)
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.2, random_state=42
    )
    # 선형 회귀 모델 학습
    lin_model = LinearRegression().fit(X_train, y_train)
    # 검증 데이터 예측
    y_pred_lin = lin_model.predict(X_test)
    # 성능 지표 계산
    results["linear"] = {
        "y_test": y_test,  # 실제값(Series; 인덱스 유지)
        "y_pred": y_pred_lin,  # 예측값(ndarray)
        "MAE": mean_absolute_error(y_test, y_pred_lin),
        "RMSE": np.sqrt(mean_squared_error(y_test, y_pred_lin)),
        "R2": r2_score(y_test, y_pred_lin),
    }

    # ---------- (B) 다중회귀 (시차 포함) ----------
    # 시차를 만들기 위해 X와 y를 통합한 임시 데이터프레임 생성
    tmp = X.copy()
    tmp = tmp.assign(Sensor=y.values)           # 종속변수 추가 (정렬/인덱스 유지)
    tmp["Sensor_t-1"] = tmp["Sensor"].shift(1)  # 1 시차
    tmp["Sensor_t-2"] = tmp["Sensor"].shift(2)  # 2 시차
    tmp = tmp.dropna()                          # 시차 생성으로 앞행 결측 발생 → 제거

    # 다중회귀 입력/타깃 구성: [원래 X 1개 + 시차 2개] = 3개 피처
    X_multi = tmp[[X.columns[0], "Sensor_t-1", "Sensor_t-2"]]  # 예: ["time_index","Sensor_t-1","Sensor_t-2"]
    y_multi = tmp["Sensor"]

    # 데이터 분할
    X_train, X_test, y_train, y_test = train_test_split(
        X_multi, y_multi, test_size=0.2, random_state=42
    )
    # 선형 회귀 학습
    multi_model = LinearRegression().fit(X_train, y_train)
    # 예측
    y_pred_multi = multi_model.predict(X_test)
    # 성능 지표
    results["multiple"] = {
        "y_test": y_test,
        "y_pred": y_pred_multi,
        "MAE": mean_absolute_error(y_test, y_pred_multi),
        "RMSE": np.sqrt(mean_squared_error(y_test, y_pred_multi)),
        "R2": r2_score(y_test, y_pred_multi),
    }

    # ---------- (C) KNN 회귀 ----------
    # 동일한 다변량 입력(X_multi)을 사용 (시차 포함)
    X_knn = X_multi.copy()
    y_knn = y_multi.copy()

    # 데이터 분할
    X_train, X_test, y_train, y_test = train_test_split(
        X_knn, y_knn, test_size=0.2, random_state=42
    )
    # K=3의 KNN 회귀 모델 (가까운 3개 이웃 평균으로 예측)
    knn_model = KNeighborsRegressor(n_neighbors=3).fit(X_train, y_train)
    # 예측
    y_pred_knn = knn_model.predict(X_test)
    # 성능 지표
    results["knn"] = {
        "y_test": y_test,
        "y_pred": y_pred_knn,
        "MAE": mean_absolute_error(y_test, y_pred_knn),
        "RMSE": np.sqrt(mean_squared_error(y_test, y_pred_knn)),
        "R2": r2_score(y_test, y_pred_knn),
    }

    # 완성된 결과 반환
    return results

# =============== [6] 두 방식별 회귀 수행 ===============
# 6-1) 시간 기반(time_index) 결과
time_results = run_regression_pipeline(df[["time_index"]], df["Sensor"])

# 6-2) 자동 순번(NUM) 기반 결과
num_results = run_regression_pipeline(df[["NUM"]], df["Sensor"])

# =============== [7] 시각화(단일 플롯 규칙 준수) ===============
def plot_result_bundle(results: dict, title: str, save_path: str):
    """
    목적 : 단일 플롯 위에 실제값 vs (선형/다중/KNN) 예측값을 함께 시각화하여 비교
    입력 :
      - results : run_regression_pipeline 반환 객체
      - title   : 그래프 제목
      - save_path : 저장 경로(파일명)
    주의 :
      - 색상/스타일 미지정(요구사항), 하나의 figure에만 구성(서브플롯 사용 금지)
    """
    plt.figure(figsize=(10, 6))  # 개별 차트 크기 지정(스타일 미적용)
    # 실제값: 단, x축은 테스트 세트의 순서(임의 샘플 순서)는 유지됨
    plt.plot(results["linear"]["y_test"].values, label="Actual", marker='o', alpha=0.7)
    # 선형 회귀 예측
    plt.plot(results["linear"]["y_pred"], label="Linear", linestyle="--")
    # 다중 회귀 예측
    plt.plot(results["multiple"]["y_pred"], label="Multiple", linestyle="-.")
    # KNN 회귀 예측
    plt.plot(results["knn"]["y_pred"], label="KNN", linestyle=":")
    # 제목/축/범례
    plt.title(title)
    plt.xlabel("테스트 샘플 인덱스")
    plt.ylabel("Sensor 값")
    plt.legend()
    # 레이아웃 정리 및 저장
    plt.tight_layout()
    plt.savefig(save_path, dpi=200)
    plt.close()

# 7-1) 시간 기반 그래프 생성
plot_result_bundle(time_results, "시간기반 회귀모형 비교", PLOT_TIME_PATH)

# 7-2) 순번 기반 그래프 생성
plot_result_bundle(num_results, "순번기반 회귀모형 비교", PLOT_NUM_PATH)

# =============== [8] Word 리포트 자동 생성 ===============
# 8-1) 문서 객체 생성
doc = Document()

# 8-2) 제목
doc.add_heading("시간기반 vs 순번기반 회귀분석 비교 리포트", 0)

# 8-3) 서론
doc.add_heading("Ⅰ. 서론", level=1)
doc.add_paragraph(
    "동일한 진동데이터(Sensor)에 대해 시간기반(time_index)과 자동순번(NUM) 기반의 회귀모형을 비교하였다. "
    "두 방식의 입력 스케일/간격 차이가 KNN·다중·단순 회귀 성능에 미치는 영향을 실증적으로 분석한다."
)

# 8-4) 결과 요약 표
doc.add_heading("Ⅱ. 분석결과 요약", level=1)
table = doc.add_table(rows=1, cols=7)
hdr = table.rows[0].cells
hdr[0].text = "모형"
hdr[1].text = "기반유형"
hdr[2].text = "MAE"
hdr[3].text = "RMSE"
hdr[4].text = "R²"
hdr[5].text = "특징"
hdr[6].text = "비고"

# 결과 행 채우기(선형/다중/KNN × 시간/순번)
for model in ["linear", "multiple", "knn"]:
    for base_name, res in [("시간", time_results), ("순번", num_results)]:
        row = table.add_row().cells
        row[0].text = model.upper()
        row[1].text = base_name
        row[2].text = f"{res[model]['MAE']:.3f}"
        row[3].text = f"{res[model]['RMSE']:.3f}"
        row[4].text = f"{res[model]['R2']:.3f}"
        row[5].text = "비선형 대응" if model == "knn" else "선형 추세 반영"
        row[6].text = "시간 스케일 영향" if (base_name == "시간" and model == "knn") else "-"

# 8-5) 시각화 삽입
doc.add_heading("Ⅲ. 시각화 비교", level=1)
doc.add_paragraph("시간기반과 순번기반의 회귀모형 예측 결과를 각각 단일 플롯으로 제시한다.")
doc.add_picture(PLOT_TIME_PATH, width=Inches(5.5))
doc.add_picture(PLOT_NUM_PATH, width=Inches(5.5))

# 8-6) 결론
doc.add_heading("Ⅳ. 결론", level=1)
doc.add_paragraph(
    "KNN 회귀는 거리 기반 알고리즘으로 입력 스케일 및 샘플 간 간격의 불균등성에 민감하다. "
    "시간기반에서는 수집 간격이 불균등할 경우 국소 밀도 차이가 커져 이웃 선택이 달라지고, "
    "순번기반은 간격이 균등하여 비교적 안정적인 결과를 보인다. "
    "따라서 시간기반 분석에는 정규화/스케일링 또는 보간을 통한 간격 균등화가 유효하다."
)

# 8-7) 참고문헌
doc.add_heading("Ⅴ. 참고문헌", level=1)
for ref in [
    "Montgomery, D. C., Jennings, C. L., & Kulahci, M. (2015). Introduction to Time Series Analysis and Forecasting. Wiley.",
    "James, G., Witten, D., Hastie, T., & Tibshirani, R. (2021). An Introduction to Statistical Learning. Springer.",
    "Hyndman, R. J., & Athanasopoulos, G. (2018). Forecasting: Principles and Practice. OTexts."
]:
    doc.add_paragraph(ref, style="List Bullet")

# 8-8) 저장
doc.save(DOCX_PATH)

# =============== [9] 실행 안내 출력 ===============
print("✅ 분석 완료")
print(f"- 시간 기반 그래프  : {PLOT_TIME_PATH}")
print(f"- 순번 기반 그래프  : {PLOT_NUM_PATH}")
print(f"- Word 리포트       : {DOCX_PATH}")
