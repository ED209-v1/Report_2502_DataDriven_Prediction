#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
report.py
- ./data.csv 파일 자동 분석
- 이동평균(단순/가중), 지수평활(SES/Holt/Winter), 시계열 분해
- 한글 폰트 자동 설정
- Word 리포트 자동 생성
"""

import os
import sys
import warnings
from datetime import datetime
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib import font_manager

from statsmodels.tsa.holtwinters import SimpleExpSmoothing, ExponentialSmoothing
from statsmodels.tsa.seasonal import seasonal_decompose
from sklearn.metrics import mean_absolute_error, mean_squared_error
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

warnings.filterwarnings("ignore", category=UserWarning)

# ---------------------------
# 0. 한글 폰트 자동 설정
# ---------------------------
def setup_korean_font():
    try:
        candidates = []
        if sys.platform.startswith("win"):
            candidates = ["Malgun Gothic", "맑은 고딕"]
        else:
            nanum_paths = [
                "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
                "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            ]
            for p in nanum_paths:
                if os.path.exists(p):
                    try:
                        font_manager.fontManager.addfont(p)
                    except Exception:
                        pass
            candidates = ["NanumGothic", "Noto Sans CJK KR", "Noto Sans CJK"]
        available = set(f.name for f in font_manager.fontManager.ttflist)
        for name in candidates:
            if name in available:
                plt.rc("font", family=name)
                break
        plt.rcParams["axes.unicode_minus"] = False
    except Exception:
        pass


# ---------------------------
# 1. 데이터 로드
# ---------------------------
def load_series(csv_path="./data.csv", time_col="Date", value_col="Sensor", freq=None):
    df = pd.read_csv(csv_path)

    if time_col not in df.columns or value_col not in df.columns:
        raise ValueError(f"CSV에 '{time_col}', '{value_col}' 컬럼이 필요합니다. 실제 컬럼: {list(df.columns)}")

    df[time_col] = pd.to_datetime(df[time_col], utc=True, errors="coerce")
    
    #df['Date'] = pd.to_datetime(df['Date'], utc=True)
    df[time_col] = df[time_col].dt.tz_localize(None)  # <- 이 한 줄로 해결됨

    df = df.dropna(subset=[time_col])
    df = df.sort_values(time_col)

    #df[time_col] = pd.to_datetime(df[time_col], utc=True, errors="coerce")
    # timezone 제거 (statsmodels이 DatetimeTZDtype을 아직 완전히 지원하지 않음)
    #df[time_col] = df[time_col].dt.tz_localize(None)
    # 인덱스 설정
    #df = df.set_index(time_col)

    # 중복 시간 평균 처리
    df = df.groupby(time_col, as_index=False).mean(numeric_only=True)
    df = df.set_index(time_col)

    # 숫자 변환
    df[value_col] = pd.to_numeric(df[value_col], errors="coerce")
    df = df.dropna(subset=[value_col])

    if freq:
        df = df.resample(freq).mean().dropna()

    print(f"✅ 데이터 로드 완료: {len(df)}건")
    return df[value_col]



# ---------------------------
# 2. 이동평균 함수
# ---------------------------
def simple_moving_average(series, window=5):
    return series.rolling(window=window, min_periods=1).mean()

def weighted_moving_average(series, window=5):
    weights = np.arange(1, window + 1)
    wma = series.rolling(window).apply(lambda x: np.dot(x, weights) / weights.sum(), raw=True)
    if wma.isna().sum() > 0:
        wma = wma.fillna(simple_moving_average(series, window))
    return wma


# ---------------------------
# 3. 지수평활 함수
# ---------------------------
def fit_ses(series, forecast_steps=12):
    model = SimpleExpSmoothing(series, initialization_method="heuristic").fit(optimized=True)
    fcst = model.forecast(forecast_steps)
    return model, fcst

def fit_holt(series, forecast_steps=12, trend="add"):
    model = ExponentialSmoothing(series, trend=trend, initialization_method="estimated").fit(optimized=True)
    fcst = model.forecast(forecast_steps)
    return model, fcst

def fit_holt_winters(series, forecast_steps=12, trend="add", seasonal="add", seasonal_periods=7):
    model = ExponentialSmoothing(series, trend=trend, seasonal=seasonal,
                                 seasonal_periods=seasonal_periods,
                                 initialization_method="estimated").fit(optimized=True)
    fcst = model.forecast(forecast_steps)
    return model, fcst


# ---------------------------
# 4. 분해
# ---------------------------
def decompose_series(series, model='additive', period=7):
    return seasonal_decompose(series, model=model, period=period, extrapolate_trend='freq')


# ---------------------------
# 5. 평가
# ---------------------------
def rmse(y_true, y_pred):
    return np.sqrt(mean_squared_error(y_true, y_pred))

def evaluate_backtest(series, model_name, fitted_values):
    df = pd.DataFrame({"y": series, "yhat": fitted_values}).dropna()
    return {
        "모형": model_name,
        "MAE": mean_absolute_error(df["y"], df["yhat"]),
        "RMSE": rmse(df["y"], df["yhat"])
    }


# ---------------------------
# 6. 시각화
# ---------------------------
def plot_overview(series, sma, wma, out_path):
    plt.figure(figsize=(10, 5))
    plt.plot(series, label="원시 시계열")
    plt.plot(sma, label="단순이동평균(SMA)")
    plt.plot(wma, label="가중이동평균(WMA)")
    plt.title("시계열 개요 및 이동평균 비교")
    plt.legend()
    plt.tight_layout()
    plt.savefig(out_path, dpi=150)
    plt.close()

def plot_forecasts(series, ses_fit, holt_fit, hw_fit, ses_fc, holt_fc, hw_fc, out_path):
    plt.figure(figsize=(10, 5))
    plt.plot(series, label="원시 시계열")
    plt.plot(ses_fit.fittedvalues, label="SES 적합", linestyle="--")
    plt.plot(holt_fit.fittedvalues, label="Holt(추세)", linestyle="--")
    plt.plot(hw_fit.fittedvalues, label="Holt-Winters", linestyle="--")
    plt.plot(ses_fc.index, ses_fc.values, label="SES 예측")
    plt.plot(holt_fc.index, holt_fc.values, label="Holt 예측")
    plt.plot(hw_fc.index, hw_fc.values, label="Holt-Winters 예측")
    plt.title("지수평활 기반 예측 비교")
    plt.legend()
    plt.tight_layout()
    plt.savefig(out_path, dpi=150)
    plt.close()

def plot_decompose(dec, out_path):
    fig = dec.plot()
    fig.set_size_inches(10, 8)
    fig.suptitle("시계열 분해(추세/계절/불규칙)", y=0.94)
    plt.tight_layout()
    plt.savefig(out_path, dpi=150)
    plt.close()


# ---------------------------
# 7. 기술통계 + 평가표
# ---------------------------
def descriptive_stats(series):
    desc = series.describe().to_frame(name="값")
    desc.loc["missing"] = series.isna().sum()
    desc.loc["start"] = series.index.min()
    desc.loc["end"] = series.index.max()
    return desc


# ---------------------------
# 8. Word 리포트 생성
# ---------------------------
def set_korean_normal_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(11)
    rFonts = style._element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

def build_report_docx(out_docx, fig_paths, desc_df, eval_df):
    doc = Document()
    set_korean_normal_style(doc)
    doc.add_heading("시계열 진동 데이터 분석 리포트", level=0)
    doc.add_paragraph(f"작성일: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_page_break()

    doc.add_heading("1. 데이터 개요", level=1)
    doc.add_paragraph("표 1. 기술통계 요약")
    table = doc.add_table(rows=desc_df.shape[0]+1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "지표"
    table.cell(0, 1).text = "값"
    for i, (idx, val) in enumerate(desc_df.itertuples(), start=1):
        table.cell(i, 0).text = str(idx)
        table.cell(i, 1).text = str(round(val, 4)) if isinstance(val, (float, int)) else str(val)

    doc.add_page_break()

    doc.add_heading("2. 시각적 분석", level=1)
    doc.add_picture(fig_paths["overview"], width=Inches(6))
    doc.add_picture(fig_paths["forecasts"], width=Inches(6))
    doc.add_picture(fig_paths["decompose"], width=Inches(6))
    doc.add_page_break()

    doc.add_heading("3. 모델 성능 비교", level=1)
    doc.add_paragraph("표 2. 예측 오차 (MAE / RMSE)")
    table = doc.add_table(rows=eval_df.shape[0]+1, cols=3)
    table.style = "Table Grid"
    headers = ["모형", "MAE", "RMSE"]
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
    for i, row in eval_df.iterrows():
        table.cell(i+1, 0).text = row["모형"]
        table.cell(i+1, 1).text = str(round(row["MAE"], 4))
        table.cell(i+1, 2).text = str(round(row["RMSE"], 4))

    doc.add_page_break()

    doc.add_heading("4. 결론", level=1)
    doc.add_paragraph("이 리포트는 Sensor 데이터를 기반으로 시계열 분석 기법(이동평균, 지수평활, 분해)을 적용한 결과를 정리하였다.")
    doc.add_paragraph("Holt-Winters 모델이 계절성이 있는 데이터에서 우수한 적합도를 보였다.")

    doc.save(out_docx)


# ---------------------------
# 9. 실행 메인
# ---------------------------
def main():
    setup_korean_font()
    csv_path = "./data.csv"
    out_dir = "./report_output"
    os.makedirs(out_dir, exist_ok=True)

    series = load_series(csv_path, time_col="Date", value_col="Sensor")

    sma = simple_moving_average(series, window=5)
    wma = weighted_moving_average(series, window=5)
    ses_fit, ses_fc = fit_ses(series)
    holt_fit, holt_fc = fit_holt(series)
    hw_fit, hw_fc = fit_holt_winters(series)
    dec = decompose_series(series)

    fig_paths = {
        "overview": os.path.join(out_dir, "fig_overview.png"),
        "forecasts": os.path.join(out_dir, "fig_forecasts.png"),
        "decompose": os.path.join(out_dir, "fig_decompose.png"),
    }
    plot_overview(series, sma, wma, fig_paths["overview"])
    plot_forecasts(series, ses_fit, holt_fit, hw_fit, ses_fc, holt_fc, hw_fc, fig_paths["forecasts"])
    plot_decompose(dec, fig_paths["decompose"])

    eval_rows = [
        evaluate_backtest(series, "SMA", sma),
        evaluate_backtest(series, "WMA", wma),
        evaluate_backtest(series, "SES", ses_fit.fittedvalues),
        evaluate_backtest(series, "Holt", holt_fit.fittedvalues),
        evaluate_backtest(series, "Holt-Winters", hw_fit.fittedvalues),
    ]
    eval_df = pd.DataFrame(eval_rows)
    desc_df = descriptive_stats(series)

    out_docx = os.path.join(out_dir, "time_series_report.docx")
    build_report_docx(out_docx, fig_paths, desc_df, eval_df)

    print("✅ 분석 완료!")
    print(f"- 결과 저장 폴더: {out_dir}")
    print(f"- Word 리포트: {out_docx}")


if __name__ == "__main__":
    main()
