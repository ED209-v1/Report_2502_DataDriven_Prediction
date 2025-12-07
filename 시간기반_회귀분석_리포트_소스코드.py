
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression
from sklearn.neighbors import KNeighborsRegressor
from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
from docx import Document
from docx.shared import Inches

# ---------------------------------
# 1. CSV 파일 로드 및 전처리
# ---------------------------------
file_path = "mixing_actuator.csv"
df = pd.read_csv(file_path)

# Date 컬럼 변환 및 정렬
df["Date"] = pd.to_datetime(df["Date"])
df = df.sort_values(by="Date").reset_index(drop=True)

# Sensor 결측 제거
df = df.dropna(subset=["Sensor"])

# 시간 차이를 초 단위로 변환하여 회귀 변수로 사용
df["time_index"] = (df["Date"] - df["Date"].min()).dt.total_seconds()

# ---------------------------------
# 2. 단순회귀분석 (시간 → Sensor)
# ---------------------------------
X = df[["time_index"]]
y = df["Sensor"]

X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

lin_model = LinearRegression()
lin_model.fit(X_train, y_train)
y_pred_lin = lin_model.predict(X_test)

lin_metrics = {
    "MAE": mean_absolute_error(y_test, y_pred_lin),
    "RMSE": np.sqrt(mean_squared_error(y_test, y_pred_lin)),
    "R2": r2_score(y_test, y_pred_lin)
}

# ---------------------------------
# 3. 다중회귀분석 (Sensor lag 변수 포함)
# ---------------------------------
df["Sensor_t-1"] = df["Sensor"].shift(1)
df["Sensor_t-2"] = df["Sensor"].shift(2)
df = df.dropna()

X_multi = df[["time_index", "Sensor_t-1", "Sensor_t-2"]]
y_multi = df["Sensor"]

X_train, X_test, y_train, y_test = train_test_split(X_multi, y_multi, test_size=0.2, random_state=42)

multi_model = LinearRegression()
multi_model.fit(X_train, y_train)
y_pred_multi = multi_model.predict(X_test)

multi_metrics = {
    "MAE": mean_absolute_error(y_test, y_pred_multi),
    "RMSE": np.sqrt(mean_squared_error(y_test, y_pred_multi)),
    "R2": r2_score(y_test, y_pred_multi)
}

# ---------------------------------
# 4. KNN 회귀분석
# ---------------------------------
X_knn = df[["time_index", "Sensor_t-1", "Sensor_t-2"]]
y_knn = df["Sensor"]

X_train, X_test, y_train, y_test = train_test_split(X_knn, y_knn, test_size=0.2, random_state=42)

knn_model = KNeighborsRegressor(n_neighbors=3)
knn_model.fit(X_train, y_train)
y_pred_knn = knn_model.predict(X_test)

knn_metrics = {
    "MAE": mean_absolute_error(y_test, y_pred_knn),
    "RMSE": np.sqrt(mean_squared_error(y_test, y_pred_knn)),
    "R2": r2_score(y_test, y_pred_knn)
}

# ---------------------------------
# 5. 결과 시각화
# ---------------------------------
plt.figure(figsize=(10, 6))
plt.plot(y_test.values, label="Actual", marker='o')
plt.plot(y_pred_lin, label="Linear", linestyle="--")
plt.plot(y_pred_multi, label="Multiple", linestyle="-.")
plt.plot(y_pred_knn, label="KNN", linestyle=":")
plt.title("예측 결과 비교 (단순, 다중, KNN 회귀)")
plt.xlabel("테스트 데이터 인덱스")
plt.ylabel("Sensor 값")
plt.legend()
plt.tight_layout()
plt.savefig("regression_comparison_timebased.png", dpi=200)
plt.close()

# ---------------------------------
# 6. Word 리포트 자동 생성
# ---------------------------------
doc = Document()
doc.add_heading("시간기반 회귀모형 비교분석 리포트", 0)

doc.add_heading("Ⅰ. 서론", level=1)
doc.add_paragraph("본 연구는 제조공정에서 수집된 진동데이터(Sensor)를 시간의 흐름에 따라 분석하여, 공정의 품질 변동을 설명할 수 있는 회귀모형을 탐색하고자 한다.")

doc.add_heading("Ⅱ. 본론", level=1)
doc.add_heading("1. 데이터 현황", level=2)
doc.add_paragraph(f"총 데이터 개수: {len(df)}개 관측치")
doc.add_paragraph(f"Sensor 평균: {df['Sensor'].mean():.3f}, 표준편차: {df['Sensor'].std():.3f}")

doc.add_heading("2. 분석 모형 설명", level=2)
doc.add_paragraph("(1) 단순회귀: 시간(time_index)과 Sensor 값의 선형 관계 추정.")
doc.add_paragraph("(2) 다중회귀: 이전 Sensor 값 포함하여 시계열 의존성 반영.")
doc.add_paragraph("(3) KNN 회귀: 인접 데이터 간 거리 기반 비선형 예측 수행.")

doc.add_heading("3. 분석 결과", level=2)
doc.add_picture("regression_comparison_timebased.png", width=Inches(6))

table = doc.add_table(rows=4, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '모델'
hdr_cells[1].text = 'MAE'
hdr_cells[2].text = 'RMSE'
hdr_cells[3].text = 'R2'

rows = [
    ("단순회귀", lin_metrics["MAE"], lin_metrics["RMSE"], lin_metrics["R2"]),
    ("다중회귀", multi_metrics["MAE"], multi_metrics["RMSE"], multi_metrics["R2"]),
    ("KNN회귀", knn_metrics["MAE"], knn_metrics["RMSE"], knn_metrics["R2"])
]

for model, mae, rmse, r2 in rows:
    row_cells = table.add_row().cells
    row_cells[0].text = model
    row_cells[1].text = f"{mae:.3f}"
    row_cells[2].text = f"{rmse:.3f}"
    row_cells[3].text = f"{r2:.3f}"

doc.add_heading("Ⅲ. 결론", level=1)
doc.add_paragraph("시간기반 회귀분석을 통해 공정 데이터의 변동성을 정량적으로 분석하였다.")

doc.save("시간기반_회귀분석_리포트_코드.docx")
